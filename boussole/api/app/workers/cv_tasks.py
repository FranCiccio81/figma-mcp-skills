"""Tâche Celery du parsing CV — ``ai.parse_cv`` (file ``ai``, F-B, D12/D16).

Chaîne (F-B nominal 3-5) :

1. ``cv_documents.status`` → ``parsing`` ; lecture du fichier en STOCKAGE
   encadrée (un échec bascule le document en ``failed``/``unreadable`` au
   lieu de le laisser bloqué en ``parsing``) ;
2. garde-fous de sûreté AVANT toute décompression
   (:mod:`app.modules.profiles.cv.safety`) : une archive DOCX au ratio de
   décompression aberrant (bombe zip : 475 Ko → 400 Mo observés) ou un PDF
   hors bornes sont refusés en ``unreadable`` — erreur PERMANENTE, aucun
   retry Celery ;
3. extraction du texte : pypdf (couche texte PDF) ou python-docx —
   PDF illisible/corrompu → ``unreadable`` ; PDF sans couche texte →
   ``image_only_pdf`` (OCR hors MVP 🟡 Q1) ; DOCX vide/corrompu →
   ``unreadable`` ; texte stocké par clé objet (``raw_text_key``, RM-B-7) ;
4. appel :func:`app.ai.tasks.extract_cv.extract_cv` (FakeProvider par
   défaut 🟡 — provider réel M5) : troncature du texte, validation stricte,
   ancrage des evidences, minimisation de la sortie, 1 retry, sinon
   ``extraction_failed`` (RM-B-3) ;
5. ``extraction_runs`` enregistré (prompt_version, model, status, sortie
   JSON validée stockée par clé objet) ; statut final ``parsed``/``failed``.

Boucle d'événements : UNE seule coroutine par tâche via ``asyncio.run``,
sur un moteur dédié ``NullPool`` (:func:`app.core.db.create_worker_engine`)
disposé en fin de coroutine — jamais le moteur global poolé.
"""

import asyncio
import io
import logging
import uuid
from collections.abc import Awaitable
from dataclasses import dataclass
from typing import Any

from sqlalchemy.ext.asyncio import async_sessionmaker

from app.ai.providers.base import LLMProvider
from app.ai.providers.fake import FakeProvider
from app.ai.tasks.extract_cv import (
    DEFAULT_MODEL,
    PROMPT_VERSION,
    CvExtraction,
    CvExtractionError,
    extract_cv,
)
from app.core.db import create_worker_engine
from app.core.storage import ObjectStorage, get_object_storage
from app.modules.profiles.cv.models import CvDocument, ExtractionRun
from app.modules.profiles.cv.safety import (
    MAX_EXTRACTED_CHARS,
    UnsafeDocumentError,
    assert_archive_safe,
    assert_pdf_pages_safe,
)
from app.modules.profiles.cv.service import DOCX_MIME, PDF_MIME
from app.workers.celery_app import celery_app

logger = logging.getLogger(__name__)


def _run[T](coro: Awaitable[T]) -> T:
    """Pont sync Celery → code async SQLAlchemy (UNE coroutine par tâche)."""
    return asyncio.run(coro)  # type: ignore[arg-type]


def _get_provider() -> LLMProvider:
    """Provider LLM de la tâche — FakeProvider 🟡 (Anthropic + fallback M5, D08)."""
    return FakeProvider()


# ------------------------------------------------------------ extraction texte


def _bounded(parts: list[str]) -> str:
    """Concatène en bornant la mémoire (:data:`MAX_EXTRACTED_CHARS`)."""
    text = "\n".join(parts)
    if len(text) > MAX_EXTRACTED_CHARS:
        return text[:MAX_EXTRACTED_CHARS]
    return text


def extract_pdf_text(data: bytes) -> str:
    """Couche texte d'un PDF via pypdf — lève sur document corrompu.

    Bornes de sûreté : nombre de pages (:data:`MAX_PDF_PAGES` 🟡) et taille
    du texte accumulé (:data:`MAX_EXTRACTED_CHARS`) — un PDF peut lui aussi
    être une bombe (des milliers de pages générées).
    """
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    assert_pdf_pages_safe(len(reader.pages))
    parts: list[str] = []
    total = 0
    for page in reader.pages:
        chunk = page.extract_text() or ""
        parts.append(chunk)
        total += len(chunk) + 1
        if total >= MAX_EXTRACTED_CHARS:
            logger.warning("cv_pdf_text_bounded chars=%d", total)
            break
    return _bounded(parts)


def extract_docx_text(data: bytes) -> str:
    """Texte d'un DOCX via python-docx (paragraphes + tableaux).

    ATTENTION : :func:`assert_archive_safe` DOIT avoir été appelée avant —
    ``python-docx`` décompresse l'archive intégralement en mémoire (une
    bombe de décompression y tue le worker par OOM).
    """
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return _bounded(parts)


def extract_text(data: bytes, mime_type: str) -> tuple[str | None, str | None]:
    """``(texte, None)`` ou ``(None, error_code)`` — fonction pure.

    Codes (enum openapi) : ``image_only_pdf`` (PDF lisible sans couche
    texte), ``unreadable`` (corrompu, vide, OU refusé par un garde-fou de
    sûreté — bombe de décompression, PDF hors bornes),
    ``unsupported_format`` (défensif — normalement rejeté en amont par le
    router, 415).

    Un refus de garde-fou est PERMANENT : il est retourné comme échec
    métier (statut ``failed``), jamais levé comme exception — la tâche
    Celery ne doit pas le rejouer.
    """
    if mime_type == PDF_MIME:
        try:
            text = extract_pdf_text(data)
        except UnsafeDocumentError as exc:
            logger.warning("cv_document_rejected mime=pdf reason=%s", exc.reason)
            return None, "unreadable"
        except Exception:
            return None, "unreadable"
        if not text.strip():
            return None, "image_only_pdf"  # scan sans couche texte (F-B alt. 1)
        return text, None
    if mime_type == DOCX_MIME:
        try:
            # AVANT toute décompression : catalogue de l'archive seulement.
            assert_archive_safe(data)
        except UnsafeDocumentError as exc:
            logger.warning("cv_document_rejected mime=docx reason=%s", exc.reason)
            return None, "unreadable"
        try:
            text = extract_docx_text(data)
        except Exception:
            return None, "unreadable"
        if not text.strip():
            return None, "unreadable"
        return text, None
    return None, "unsupported_format"


# ------------------------------------------------------------ orchestration


@dataclass(frozen=True, slots=True)
class ParseOutcome:
    """Résultat pur du parsing d'un document (avant persistance)."""

    document_status: str  # 'parsed' | 'failed'
    error_code: str | None
    run_status: str | None  # 'success' | 'schema_error' | 'failed' | None (pas d'appel LLM)
    raw_text: str | None
    extraction: CvExtraction | None


def run_cv_extraction(data: bytes, mime_type: str, provider: LLMProvider) -> ParseOutcome:
    """Extraction texte + appel ``extract_cv`` — fonction pure (testable sans DB)."""
    text, error_code = extract_text(data, mime_type)
    if text is None:
        return ParseOutcome(
            document_status="failed",
            error_code=error_code,
            run_status=None,  # aucun appel LLM : pas de run (R7)
            raw_text=None,
            extraction=None,
        )
    try:
        extraction = extract_cv(text, provider)
    except CvExtractionError:
        # RM-B-3 : retry + échec propre déjà joués dans extract_cv.
        return ParseOutcome(
            document_status="failed",
            error_code="extraction_failed",
            run_status="schema_error",
            raw_text=text,
            extraction=None,
        )
    except Exception:
        logger.exception("cv_extraction_provider_error")
        return ParseOutcome(
            document_status="failed",
            error_code="extraction_failed",
            run_status="failed",
            raw_text=text,
            extraction=None,
        )
    return ParseOutcome(
        document_status="parsed",
        error_code=None,
        run_status="success",
        raw_text=text,
        extraction=extraction,
    )


async def _parse_cv_cycle(
    cv_document_id: str,
    *,
    storage: ObjectStorage | None = None,
    provider: LLMProvider | None = None,
) -> dict[str, Any]:
    """Cycle complet dans UNE coroutine : moteur NullPool dédié, disposé en fin."""
    storage = storage if storage is not None else get_object_storage()
    provider = provider if provider is not None else _get_provider()
    engine = create_worker_engine()
    try:
        factory = async_sessionmaker(engine, expire_on_commit=False)
        async with factory() as session:
            document = await session.get(CvDocument, uuid.UUID(cv_document_id))
            if document is None:
                logger.warning("parse_cv_document_missing id=%s", cv_document_id)
                return {"status": "missing"}
            document.status = "parsing"
            document.error_code = None
            await session.commit()  # statut visible pendant le parsing (polling front)

            try:
                data = storage.get(document.file_key)
            except Exception:
                # Sans ce garde-fou, une lecture de stockage en échec laisse
                # le document en « parsing » indéfiniment (l'exception
                # remonte à Celery, qui rejoue, et le statut n'est jamais
                # rétabli) — échec propre et visible à la place.
                logger.exception("cv_storage_read_failed id=%s", cv_document_id)
                document.status = "failed"
                document.error_code = "unreadable"
                await session.commit()
                return {"status": "failed", "error_code": "unreadable"}

            outcome = run_cv_extraction(data, document.mime_type, provider)

            if outcome.raw_text is not None:
                raw_text_key = f"cv/{document.user_id}/{document.id}/raw_text.txt"
                storage.put(raw_text_key, outcome.raw_text.encode("utf-8"))
                document.raw_text_key = raw_text_key

            if outcome.run_status is not None:
                run = ExtractionRun(
                    id=uuid.uuid4(),
                    cv_document_id=document.id,
                    prompt_version=PROMPT_VERSION,
                    model=DEFAULT_MODEL,
                    status=outcome.run_status,
                    output_key=None,
                )
                if outcome.extraction is not None:
                    output_key = (
                        f"cv/{document.user_id}/{document.id}/extraction/{run.id}.json"
                    )
                    storage.put(
                        output_key, outcome.extraction.model_dump_json().encode("utf-8")
                    )
                    run.output_key = output_key
                session.add(run)

            document.status = outcome.document_status
            document.error_code = outcome.error_code
            await session.commit()
            return {"status": document.status, "error_code": document.error_code}
    finally:
        await engine.dispose()


@celery_app.task(
    name="ai.parse_cv",
    bind=True,
    max_retries=3,  # retries infra (DB/stockage indisponible) — D08 🟡
    autoretry_for=(Exception,),
    retry_backoff=10,
    retry_backoff_max=300,
    retry_jitter=True,
)
def parse_cv(self: Any, cv_document_id: str) -> dict[str, Any]:
    """Parsing asynchrone d'un CV importé (F-B) — file ``ai``.

    Les échecs MÉTIER (illisible, PDF image, extraction invalide) sont
    persistés sur la ressource (``status='failed'`` + ``error_code``) et ne
    déclenchent PAS de retry Celery ; seuls les incidents d'infrastructure
    (exceptions) sont rejoués avec backoff exponentiel.
    """
    report = _run(_parse_cv_cycle(cv_document_id))
    logger.info("parse_cv_done id=%s report=%s", cv_document_id, report)
    return report
